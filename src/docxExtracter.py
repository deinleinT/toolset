'''
Created on Apr 18, 2018

@author: deinlein thomas
'''
    
import PyPDF2 
import os
import sys
import threading
import docx


class MyThread(threading.Thread):

    def __init__(self, docxList, wholeText):
        threading.Thread.__init__(self)
        self.__docxList = docxList
        self.__wholeText = wholeText
    
    def run(self):
        print("")
        print("Starting Thread " + str(threading.current_thread().getName()) + " Number of Docxs: " + str(len(self.__docxList)))
        counterProgress = 0
        for s in self.__docxList: 
            progress = (int) (counterProgress / len(self.__docxList) * 100)
            print("Processing Specification " + str(s) + " in Thread " + str(threading.current_thread().getName()) + ", Thread progress: " + str(progress) + "% ...\n")
            text = ""
            doc = docx.Document(s)
            try:
                if wholeText:
                    for para in doc.paragraphs:
                        text += str(para.text + "\n")
                else:
                    print("")
                    scope = "Scope"
                    scopeCount = 0
                    references = "References"
                    # just scope
                    for para in doc.paragraphs:
                        if scope in para.text and scopeCount == 0:
                            scopeCount += 1
                            continue
                        elif scope in para.text and scopeCount == 1:
                            scopeCount = 2
                            text += str(para.text + "\n")
                        elif scopeCount == 2 and references not in para.text:
                            text += str(para.text + "\n")
                        elif references in para.text and scopeCount == 2:
                            break
                    
                allScopes[s] = text
            except Exception as e:
                print (str(e))
                allScopes[s] = "Error with Specification " + str(s) + " " + str(e)
                continue
            
            try:
                title = str(doc.core_properties.subject)
            except Exception as e:
                title = "Was unreadable from DOCX!"
            allTitles[s] = title
            counterProgress += 1
        print("\n\n*** Finished Thread " + str(threading.current_thread().getName()) + "\n\n")
    

def get_arg(index):
    return sys.argv[index]


#########################################################################
try:
    start = get_arg(1)
    stop = get_arg(2)
    textFileName = get_arg(3)
    direct = get_arg(4)
except IndexError as e:
    start = "scope"
    stop = "references"
    textFileName = "scopesOfSpecifications"
    direct = "./Specifications"

#############
wholeText = False
if "-a" in sys.argv:
    wholeText = True
if wholeText:
    textFileName = "textOfAllSpecifications"
###############

dirList = os.listdir(direct)
dirList.sort()
pdfs = []

for g in dirList:
    if ".docx" in g:
        pdfs.append(direct + "/" + g)

numProcessors = (int)(os.environ['NUMBER_OF_PROCESSORS']) * 2

allScopes = dict()
allTitles = dict()
###############################
# 
threads = []
print("Starting with reading Docx...\n")
print("Number of all DOCX: " + str(len(pdfs)))

startTwo = 0
intervalTwo = round(len(pdfs) / numProcessors) - 1
if intervalTwo <= 0:
    intervalTwo = 1;

# the first thread of each excel must pass True in constructor
threadOne = MyThread(pdfs[startTwo:intervalTwo], wholeText)
startTwo += intervalTwo

# append all to array
threads.append(threadOne)

while startTwo < len(pdfs):
    endTwo = startTwo + intervalTwo + 1
    if endTwo > len(pdfs):
        endTwo = len(pdfs)
    threads.append(MyThread(pdfs[startTwo:endTwo], wholeText))  
    startTwo = endTwo 

# join
for t in threads:
    t.start()

for t in threads:
    t.join()
print("\n*******Threads finished!*********\n")

##############################
######

# write into textfile
print("Write the file...\n")
if textFileName != "":
    out = open(textFileName + ".txt", "w")
else:
    out = open("allScopes.txt", "w")
out.write("\n\n")
out.write("-------------------------------------------------------------------------------------------------------------------------------------\n\n")
for s in sorted(allScopes):
    out.write("PDF: " + s + "\n\n")
    out.write("Title: " + str(allTitles.get(s, " *** ERROR OCCURED ***")) + "\n\n")
    out.write("All Text from " + start + ":\n")
    try:
        out.write(str(allScopes[s]))
    except (UnicodeEncodeError, UnicodeDecodeError, UnicodeError):
        out.write("ERROR IN THIS Specification")
    out.write("\n\n")
    out.write("-------------------------------------------------------------------------------------------------------------------------------------\n\n")
out.close()
        
print("\n\n\n FINISHED!")
